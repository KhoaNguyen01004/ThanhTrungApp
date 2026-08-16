#!/usr/bin/env python3
"""Build a K17 internship report (.docx) from a Markdown draft.

Usage:
    python build_report_docx.py draft.md -o "Bao_cao_thuc_tap.docx"
    python build_report_docx.py draft.md -o out.docx --pagecount   # needs soffice

Output format follows "HUONG DAN VIET BAO CAO THUC TAP THUC TE" (K17):
A4, Times New Roman 13, line spacing 1.5, margins T2 / B2 / L3.5 / R2 cm,
page number centered in the header, lowercase roman for the front matter and
arabic from MO DAU onward.

Input Markdown layout
---------------------
    ---
    truong: Dai hoc Quoc te Sai Gon (SIU)
    khoa: Khoa Khoa hoc May tinh
    nganh: Khoa hoc May tinh
    de_tai: ...
    mssv: 1234567890
    sinh_vien: Nguyen A
    gvhd: Nguyen C
    don_vi: ...
    thang: 08
    nam: 2026
    logo: assets/logo.png        # optional
    ---

    # LOI CAM ON
    <text>

    # MO DAU
    ## 1. Ly do chon de tai
    ...

    # CHUONG 1: GIOI THIEU TONG QUAN
    ## 1.1. Tong quan co so ly thuyet
    ### 1.1.1. ...

Everything before the `# MO DAU` heading is front matter (roman numerals).
Missing front-matter pages (nhan xet, muc luc, danh muc) are inserted
automatically as blank signable pages.

Supported Markdown: headings h1-h4, paragraphs, `-`/`*` bullets, `1.` numbered
lists, pipe tables, `![caption](path)` images, `**bold**`, `*italic*`,
`> ` blockquotes (rendered as indented notes). Fenced code blocks render in
Consolas 11 with no first-line indent.
"""

from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Times New Roman"
BODY_PT = 13
MONO = "Consolas"

# Front-matter pages required by the guideline, in order, with their headings.
FRONT_MATTER = [
    ("LOI CAM ON", "LỜI CẢM ƠN", "signature_student"),
    ("NHAN XET CUA DON VI THUC TAP", "NHẬN XÉT CỦA ĐƠN VỊ THỰC TẬP", "signature_company"),
    ("NHAN XET CUA GIANG VIEN HUONG DAN", "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", "signature_lecturer"),
    ("MUC LUC", "MỤC LỤC", "toc"),
    ("DANH MUC CAC CHU VIET TAT", "DANH MỤC CÁC CHỮ VIẾT TẮT", "list"),
    ("DANH MUC HINH ANH, SO DO, BANG BIEU", "DANH MỤC HÌNH ẢNH, SƠ ĐỒ, BẢNG BIỂU", "list"),
]

DOTTED = "." * 95


# ---------------------------------------------------------------- utilities

def strip_accents(s: str) -> str:
    import unicodedata

    return "".join(
        c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn"
    ).replace("đ", "d").replace("Đ", "D")


def key_of(heading: str) -> str:
    return re.sub(r"[^A-Z ]", "", strip_accents(heading).upper()).strip()


def set_run_font(run, size=BODY_PT, bold=False, italic=False, name=FONT, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def add_field(paragraph, instr: str, placeholder: str = "", color=None):
    """Insert a Word field code (PAGE, TOC, ...) into a paragraph.

    The colour is set explicitly: LibreOffice otherwise renders a PAGE field
    result using colour picked up elsewhere in the section, which turned the
    body page numbers red once a red run existed in that section.
    """
    color = color or RGBColor(0, 0, 0)

    # Three runs, not one: the field result must be its own run carrying its
    # own rPr, or LibreOffice formats the result with whatever run formatting
    # it last saw in the section (a highlighted run elsewhere in the body was
    # enough to highlight every page number).
    r1 = paragraph.add_run()
    set_run_font(r1, color=color)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r1._r.append(begin)
    r1._r.append(instr_el)

    r2 = paragraph.add_run()
    set_run_font(r2, color=color)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r2._r.append(sep)

    r3 = paragraph.add_run(placeholder or "")
    set_run_font(r3, color=color)

    r4 = paragraph.add_run()
    set_run_font(r4, color=color)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r4._r.append(end)


def set_page_numbering(section, fmt: str | None, start: int | None):
    """fmt: 'lowerRoman' | 'decimal' | None (inherit)."""
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(old)
    if fmt is None and start is None:
        return
    el = OxmlElement("w:pgNumType")
    if fmt:
        el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))
    anchor = sect_pr.find(qn("w:pgMar"))
    if anchor is not None:
        anchor.addnext(el)
    else:
        sect_pr.append(el)


def header_page_number(section, enabled: bool):
    """Guideline: page number centered at the TOP of every page."""
    section.header.is_linked_to_previous = False
    hdr = section.header
    for p in list(hdr.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    p = hdr.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if enabled:
        add_field(p, " PAGE ", "1")


def configure_section(section, numbering_fmt, start, show_number=True):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    set_page_numbering(section, numbering_fmt, start)
    header_page_number(section, show_number)


def enable_update_fields(doc):
    """Make Word refresh the TOC when the file is opened."""
    settings = doc.settings.element
    for old in settings.findall(qn("w:updateFields")):
        settings.remove(old)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


# ------------------------------------------------------------------ styling

def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading styles carry the TOC levels, so they must stay Heading 1/2/3.
    sizes = {1: 14, 2: 13, 3: 13, 4: 13}
    theme_attrs = ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme")
    for level, size in sizes.items():
        # The linked character style keeps theme fonts and the blue accent
        # colour, which win over the paragraph style in Word and LibreOffice.
        char_id = f"Heading {level} Char"
        if char_id in [s.name for s in doc.styles]:
            crpr = doc.styles[char_id].element.get_or_add_rPr()
            crf = crpr.find(qn("w:rFonts"))
            if crf is not None:
                for attr in theme_attrs:
                    crf.attrib.pop(qn(attr), None)
                for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    crf.set(qn(attr), FONT)
            for tag in ("w:color", "w:sz", "w:szCs"):
                for el in crpr.findall(qn(tag)):
                    crpr.remove(el)
        st = doc.styles[f"Heading {level}"]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.all_caps = False
        srpr = st.element.get_or_add_rPr()
        srf = srpr.find(qn("w:rFonts"))
        if srf is None:
            srf = OxmlElement("w:rFonts")
            srpr.insert(0, srf)
        for attr in theme_attrs:
            srf.attrib.pop(qn(attr), None)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            srf.set(qn(attr), FONT)
        spf = st.paragraph_format
        spf.line_spacing = 1.5
        spf.space_before = Pt(12 if level == 1 else 6)
        spf.space_after = Pt(6)
        spf.keep_with_next = True
        spf.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        )


# ------------------------------------------------------------- md rendering

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_inline(paragraph, text: str, size=BODY_PT, base_bold=False, base_italic=False):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = paragraph.add_run(part[2:-2])
            set_run_font(r, size, bold=True, italic=base_italic)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            set_run_font(r, size, bold=base_bold, italic=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            set_run_font(r, size - 2, name=MONO)
        else:
            r = paragraph.add_run(part)
            set_run_font(r, size, bold=base_bold, italic=base_italic)


def body_paragraph(doc, text, indent_first=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.0)
    add_inline(p, text)
    return p


def add_placeholder_box(doc, caption: str, height_cm: float = 7.0):
    """A bordered box standing in for a figure the author will draw later.

    Rendered instead of silently dropping an image whose file does not exist,
    so the gap is visible in the printed draft and impossible to forget.
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(14)
    tr = table.rows[0]._tr
    trpr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trpr.append(h)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(int(height_cm * 28.35 / 3))
    set_run_font(
        p.add_run(f"[Chỗ dành cho hình — {caption}]" if caption else "[Chỗ dành cho hình]"),
        12,
        italic=True,
        color=RGBColor(0x80, 0x80, 0x80),
    )
    return table


def add_chapter_heading(doc, text):
    """Level-1 heading, centered, explicit run font so no theme font leaks in."""
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), 14, bold=True)
    return p


def blank_lines(doc, n, size=BODY_PT):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(""), size)


def centered(doc, text, size=BODY_PT, bold=False, italic=False, spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), size, bold=bold, italic=italic)
    return p


# ------------------------------------------------------------------ parsing

def parse_front_matter(text: str):
    meta = {}
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            for line in text[3:end].strip().splitlines():
                if ":" in line and not line.strip().startswith("#"):
                    k, v = line.split(":", 1)
                    meta[k.strip()] = v.strip().strip('"').strip("'")
            text = text[end + 4 :]
    return meta, text.lstrip("\n")


def _split_on(md: str, level: int):
    marker = "#" * level + " "
    deeper = "#" * (level + 1)
    sections, current, buf = [], None, []
    for line in md.splitlines():
        if line.startswith(marker) and not line.startswith(deeper):
            if current is not None:
                sections.append((current, buf))
            current, buf = line[len(marker) :].strip(), []
        elif current is not None:
            # promote sub-headings so the caller always sees h2/h3/h4
            m = re.match(r"^(#{2,6})(\s+.*)$", line)
            if level > 1 and m:
                buf.append("#" * max(2, len(m.group(1)) - level + 1) + m.group(2))
            else:
                buf.append(line)
        else:
            buf.append(line)
    if current is not None:
        sections.append((current, buf))
    return sections


def split_sections(md: str):
    """Split on top-level headings -> [(heading, body_lines)].

    Drafts written with `##` as the top level (no `#` at all) are handled by
    falling back one level and promoting everything beneath it.
    """
    sections = _split_on(md, 1)
    if len(sections) < 2:
        fallback = _split_on(md, 2)
        if len(fallback) > len(sections):
            return fallback
    return sections


# ------------------------------------------------------------- page builders

def build_cover(doc, meta):
    logo = meta.get("logo")
    blank_lines(doc, 1)
    if logo and os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Cm(3.5))
        blank_lines(doc, 1)
    centered(doc, meta.get("truong", "").upper(), 14, bold=True)
    if meta.get("khoa"):
        centered(doc, meta["khoa"].upper(), 13, bold=True)
    blank_lines(doc, 4)
    centered(doc, "BÁO CÁO", 20, bold=True)
    centered(doc, "THỰC TẬP THỰC TẾ", 20, bold=True)
    blank_lines(doc, 1)
    if meta.get("de_tai"):
        centered(doc, "Đề tài:", 13, italic=True)
        centered(doc, meta["de_tai"].upper(), 15, bold=True)
    blank_lines(doc, 4)
    for label, key in (
        ("Ngành:", "nganh"),
        ("MSSV:", "mssv"),
        ("Họ tên sinh viên:", "sinh_vien"),
        ("Giảng viên hướng dẫn:", "gvhd"),
        ("Đơn vị thực tập:", "don_vi"),
    ):
        if meta.get(key):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(f"{label} "), 13)
            set_run_font(p.add_run(meta[key]), 13, bold=True)
    blank_lines(doc, 4)
    centered(doc, f"Thành phố Hồ Chí Minh – {meta.get('nam', '')}", 13, bold=True)


def build_signature_block(doc, meta, kind):
    blank_lines(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run(
            f"Tp. Hồ Chí Minh, ngày    tháng {meta.get('thang', '08')} năm {meta.get('nam', '')}"
        ),
        13,
        italic=True,
    )
    lines = {
        "signature_student": ["Sinh viên thực hiện", "(Ký, ghi rõ họ tên)"],
        "signature_company": ["(Ký, ghi rõ họ tên, chức vụ và đóng dấu)"],
        "signature_lecturer": ["Giảng viên hướng dẫn", "(Ký, ghi rõ họ tên)"],
    }[kind]
    for i, line in enumerate(lines):
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        q.paragraph_format.line_spacing = 1.5
        q.paragraph_format.space_after = Pt(0)
        set_run_font(q.add_run(line), 13, bold=(i == 0 and len(lines) > 1))


def build_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(
        p,
        ' TOC \\o "1-3" \\h \\z \\u ',
        "Nhấn Ctrl+A rồi F9 trong Word để cập nhật mục lục.",
    )


def build_front_page(doc, meta, heading_vi, kind, content_lines, render_body):
    add_chapter_heading(doc, heading_vi)
    blank_lines(doc, 1)
    if kind == "toc":
        build_toc(doc)
        return
    if content_lines and any(l.strip() for l in content_lines):
        render_body(doc, content_lines)
        if kind.startswith("signature"):
            build_signature_block(doc, meta, kind)
        return
    if kind.startswith("signature"):
        for _ in range(18):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(DOTTED), 13)
        build_signature_block(doc, meta, kind)


# ------------------------------------------------------------ body rendering

def render_body(doc, lines):
    i = 0
    fig_no = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank lines and markdown horizontal rules carry no meaning in Word
        if not stripped or re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            i += 1
            continue

        # HTML comments hold data provenance notes; they stay in the .md only
        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # missing-data marker: [[CAN SO LIEU: mo ta]] -> visible red note
        m = re.match(r"^\[\[\s*(?:CAN SO LIEU|CẦN SỐ LIỆU|TODO)\s*:\s*(.*?)\]\]$", stripped, re.I)
        if m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            # Plain bold, deliberately: LibreOffice renders a section's PAGE
            # field result with the colour/highlight of other runs in that
            # section, so a red or highlighted marker also paints every page
            # number. Brackets and caps make it findable without formatting.
            set_run_font(p.add_run(f"[CẦN SỐ LIỆU: {m.group(1).strip()}]"), 13, bold=True)
            i += 1
            continue

        # fenced code
        if stripped.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            for c in code:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1.0)
                set_run_font(p.add_run(c), 11, name=MONO)
            continue

        # headings h2-h4
        m = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, size=13 if level > 1 else 14, base_bold=True)
            for r in p.runs:
                r.bold = True
            i += 1
            continue

        # image, or a placeholder box when the file is not there yet
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if m:
            caption, path = m.group(1), m.group(2)
            if path and os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.add_run().add_picture(path, width=Cm(14))
            else:
                add_placeholder_box(doc, caption)
            fig_no += 1
            if caption:
                centered(doc, caption, 12, italic=True, spacing=1.0)
            blank_lines(doc, 1)
            i += 1
            continue

        # explicit placeholder: [[CHO TRONG: Hình 3.1. Sơ đồ quy trình]]
        m = re.match(r"^\[\[\s*(?:CHO TRONG|CHỖ TRỐNG|PLACEHOLDER)\s*:\s*(.*?)\]\]$", stripped, re.I)
        if m:
            caption = m.group(1).strip()
            add_placeholder_box(doc, caption)
            if caption:
                centered(doc, caption, 12, italic=True, spacing=1.0)
            blank_lines(doc, 1)
            i += 1
            continue

        # pipe table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:\-|]+\|$", lines[i + 1].strip()
        ):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(
                    [c.strip() for c in lines[i].strip().strip("|").split("|")]
                )
                i += 1
            header, body = rows[0], rows[2:]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(cell.paragraphs[0], text, size=12, base_bold=True)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in body:
                cells = table.add_row().cells
                for c, text in enumerate(row[: len(header)]):
                    add_inline(cells[c].paragraphs[0], text, size=12)
            continue

        # blockquote
        if stripped.startswith(">"):
            p = body_paragraph(doc, stripped.lstrip("> ").strip(), indent_first=False)
            p.paragraph_format.left_indent = Cm(1.0)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # bullets / numbers
        m = re.match(r"^([-*+])\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, m.group(1))
            i += 1
            continue

        # paragraph: join wrapped lines
        buf = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,4}\s|[-*+]\s|\d+[.)]\s|\||>|!\[|```)", lines[i].strip()
        ):
            buf.append(lines[i].strip())
            i += 1
        body_paragraph(doc, " ".join(buf))


# ---------------------------------------------------------------------- main

def build(md_path: str, out_path: str) -> str:
    raw = open(md_path, encoding="utf-8").read()
    meta, md = parse_front_matter(raw)
    sections = split_sections(md)

    by_key = {key_of(h): (h, b) for h, b in sections}
    body_start = None
    for idx, (h, _) in enumerate(sections):
        if key_of(h).startswith("MO DAU"):
            body_start = idx
            break
    if body_start is None:
        body_start = 0
        print("warning: no '# MỞ ĐẦU' heading found — everything treated as body")

    front_sections = sections[:body_start]
    body_sections = sections[body_start:]
    front_keys = {key_of(h) for h, _ in front_sections}

    doc = Document()
    configure_styles(doc)

    # --- section 1: cover, no page number
    configure_section(doc.sections[0], None, 1, show_number=False)
    build_cover(doc, meta)

    # --- section 2: front matter, lowercase roman
    doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(doc.sections[1], "lowerRoman", 1, show_number=True)

    first = True
    for k, heading_vi, kind in FRONT_MATTER:
        if not first:
            doc.add_page_break()
        first = False
        content = by_key.get(k, (None, []))[1] if k in front_keys else []
        build_front_page(doc, meta, heading_vi, kind, content, render_body)
    # any extra front-matter section the author added
    for h, b in front_sections:
        if key_of(h) not in {k for k, _, _ in FRONT_MATTER}:
            doc.add_page_break()
            add_chapter_heading(doc, h)
            render_body(doc, b)

    # --- section 3: body, arabic restarting at 1
    doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(doc.sections[2], "decimal", 1, show_number=True)

    for n, (heading, lines) in enumerate(body_sections):
        if n:
            doc.add_page_break()
        add_chapter_heading(doc, heading)
        render_body(doc, lines)

    enable_update_fields(doc)
    doc.save(out_path)
    return out_path


def page_count(docx_path: str) -> int | None:
    """Render with LibreOffice to count real pages (guideline: body >= 30)."""
    with tempfile.TemporaryDirectory() as td:
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", td, docx_path],
                check=True,
                capture_output=True,
                timeout=300,
            )
        except (OSError, subprocess.SubprocessError):
            return None
        pdf = os.path.join(
            td, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        )
        if not os.path.exists(pdf):
            return None
        data = open(pdf, "rb").read()
        return data.count(b"/Type /Page") - data.count(b"/Type /Pages") or len(
            re.findall(rb"/Type\s*/Page[^s]", data)
        )


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("markdown")
    ap.add_argument("-o", "--output", required=True)
    ap.add_argument(
        "--pagecount",
        action="store_true",
        help="render via LibreOffice and report the page count",
    )
    args = ap.parse_args()

    out = build(args.markdown, args.output)
    print(f"wrote {out}")
    if args.pagecount:
        n = page_count(out)
        if n is None:
            print("page count unavailable (LibreOffice not found)")
        else:
            print(f"pages: {n} (guideline minimum for the body: 30)")
            if n < 33:
                print("  -> body likely under 30 pages; expand chapters 2-4")


if __name__ == "__main__":
    sys.exit(main())
